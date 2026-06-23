#!/usr/bin/env python3
"""
Convert a Projects or Accelerators source file (.json / .docx / .xlsx) into the
canonical JSON shape consumed by the commercetools Insights deploy agent.

Usage:
    python3 data_to_json.py --kind projects     <input.{json,docx,xlsx}>
    python3 data_to_json.py --kind accelerators <input.{json,docx,xlsx}>

Prints canonical JSON to stdout. Exit non-zero on a fatal parse error.

Canonical shapes (must match data/projects.ts and data/accelerators.ts):
  projects:     { org, generated, projects[] { id, name, status, members[]{name,designation} } }
  accelerators: { org, title, heading, intro, generated,
                  items[] { n, slug, name, category, tagline, maturity, customerValue,
                            overview, benefits[{title,text}], technical[], challenge{title,text},
                            resolves[], idealCustomer, signals[], valueCommercetools, valueRoyalCyber } }

Field labels/headers are matched case- and punctuation-insensitively, so both
human labels ("Customer Value", "Value for commercetools", "Challenge Title") and
the canonical camelCase keys ("customerValue", "valueCommercetools", "challengeTitle")
are accepted.
"""
import sys, os, json, re, argparse, datetime

ORG_DEFAULT = "Royal Cyber"
ACC_TITLE_DEFAULT = "Royal Cyber Accelerators"
ACC_HEADING_DEFAULT = "Four Accelerators, Built for Customer Value"
ACC_INTRO_DEFAULT = ("Each accelerator extends commercetools with applied AI, analytics and automation "
                     "— and answers a real customer pain. Built and proven by the Royal Cyber "
                     "commercetools Center of Excellence, they shorten time-to-value across B2B commerce, "
                     "platform migration, analytics and catalog management.")

def today():
    return datetime.date.today().strftime("%b %d, %Y")

def clean(s):
    if s is None:
        return ""
    return re.sub(r"\s+", " ", str(s).replace("\xa0", " ")).strip()

def lines(cell):
    if cell is None:
        return []
    return [clean(p) for p in re.split(r"[\r\n]+", str(cell)) if clean(p)]

def slugify(s):
    s = re.sub(r"[^a-z0-9]+", "-", clean(s).lower())
    return re.sub(r"(^-+|-+$)", "", s)

def norm_key(s):
    return re.sub(r"[^a-z0-9]", "", clean(s).lower())

# ----------------------------------------------------------------------------- PROJECTS
PROJ_PROJECT = {"projectname", "project"}
PROJ_NAME = {"employeename", "name", "member", "teammember"}
PROJ_DESIG = {"designation", "role"}
PROJ_STATUS = {"status"}

def projects_from_docx(path):
    from docx import Document
    doc = Document(path)
    projects = []
    for t in doc.tables:
        rows = t.rows
        if not rows:
            continue
        head0 = clean(rows[0].cells[0].text)
        m = re.search(r"project\s*name\s*:?\s*(.*)$", head0, re.I)
        name = clean(m.group(1)) if (m and m.group(1)) else head0
        if not name:
            continue
        name_col, desig_col, hdr_idx = 1, 2, None
        for i, r in enumerate(rows[1:], start=1):
            keys = [norm_key(c.text) for c in r.cells]
            if any(k in PROJ_NAME for k in keys):
                hdr_idx = i
                for j, k in enumerate(keys):
                    if k in PROJ_NAME:
                        name_col = j
                    if k in PROJ_DESIG:
                        desig_col = j
                break
        if desig_col == name_col:
            desig_col = name_col + 1
        members = []
        start = (hdr_idx + 1) if hdr_idx is not None else 1
        for r in rows[start:]:
            cells = r.cells
            nm = clean(cells[name_col].text) if len(cells) > name_col else ""
            dg = clean(cells[desig_col].text) if len(cells) > desig_col else ""
            if not nm or norm_key(nm) in PROJ_NAME:
                continue
            members.append({"name": nm, "designation": dg})
        if members:
            projects.append({"id": slugify(name), "name": name, "status": "current", "members": members})
    return {"org": ORG_DEFAULT, "generated": today(), "projects": projects}

def projects_from_xlsx(path):
    import openpyxl
    wb = openpyxl.load_workbook(path, data_only=True)
    ws = wb.active
    rows = [list(r) for r in ws.iter_rows(values_only=True)]
    rows = [r for r in rows if any(clean(c) for c in r)]
    if not rows:
        return {"org": ORG_DEFAULT, "generated": today(), "projects": []}
    header = [norm_key(c) for c in rows[0]]
    def col(group):
        for i, h in enumerate(header):
            if h in group:
                return i
        return None
    pc, nc, dc, sc = col(PROJ_PROJECT), col(PROJ_NAME), col(PROJ_DESIG), col(PROJ_STATUS)
    if pc is None or nc is None:
        raise ValueError("projects.xlsx must have at least 'Project Name' and 'Employee Name' columns")
    order, pmap = [], {}
    for r in rows[1:]:
        pname = clean(r[pc]) if pc < len(r) else ""
        nm = clean(r[nc]) if nc < len(r) else ""
        if not pname or not nm:
            continue
        dg = clean(r[dc]) if dc is not None and dc < len(r) else ""
        st = clean(r[sc]).lower() if sc is not None and sc < len(r) else ""
        if pname not in pmap:
            pmap[pname] = {"id": slugify(pname), "name": pname,
                           "status": st if st in ("current", "past") else "current", "members": []}
            order.append(pname)
        if st in ("current", "past"):
            pmap[pname]["status"] = st
        pmap[pname]["members"].append({"name": nm, "designation": dg})
    return {"org": ORG_DEFAULT, "generated": today(), "projects": [pmap[p] for p in order]}

# ------------------------------------------------------------------------- ACCELERATORS
ACC_FIELD_MAP = {
    "n": "n", "slug": "slug", "name": "name", "category": "category", "tagline": "tagline",
    "maturity": "maturity", "customervalue": "customerValue", "overview": "overview",
    "challengetitle": "challengeTitle", "challenge": "challengeTitle", "challengetext": "challengeText",
    "benefits": "benefits", "technical": "technical", "technicalstandpoint": "technical",
    "resolves": "resolves", "howitresolvesit": "resolves", "howthisacceleratorresolvesit": "resolves",
    "idealcustomer": "idealCustomer", "idealcustomerprofile": "idealCustomer",
    "signals": "signals", "signalstolookfor": "signals",
    "valuecommercetools": "valueCommercetools", "valueforcommercetools": "valueCommercetools",
    "valueroyalcyber": "valueRoyalCyber", "valueforroyalcyber": "valueRoyalCyber",
}
META_KEYS = {"org": "org", "title": "title", "heading": "heading", "intro": "intro", "generated": "generated"}

def _benefits(raw):
    out = []
    for ln in (raw if isinstance(raw, list) else lines(raw)):
        if isinstance(ln, dict):
            out.append({"title": clean(ln.get("title")), "text": clean(ln.get("text"))}); continue
        if "::" in ln:
            t, d = ln.split("::", 1)
        elif " - " in ln:
            t, d = ln.split(" - ", 1)
        else:
            t, d = ln, ""
        out.append({"title": clean(t), "text": clean(d)})
    return out

def _as_list(v):
    return v if isinstance(v, list) else lines(v)

def _build_item(d, idx):
    name = clean(d.get("name", ""))
    return {
        "n": clean(d.get("n")) or f"{idx:02d}",
        "slug": clean(d.get("slug")) or slugify(name),
        "name": name,
        "category": clean(d.get("category", "")),
        "tagline": clean(d.get("tagline", "")),
        "maturity": clean(d.get("maturity", "")),
        "customerValue": clean(d.get("customerValue", "")),
        "overview": clean(d.get("overview", "")),
        "benefits": _benefits(d.get("benefits", [])),
        "technical": _as_list(d.get("technical")),
        "challenge": {"title": clean(d.get("challengeTitle", "")), "text": clean(d.get("challengeText", ""))},
        "resolves": _as_list(d.get("resolves")),
        "idealCustomer": clean(d.get("idealCustomer", "")),
        "signals": _as_list(d.get("signals")),
        "valueCommercetools": clean(d.get("valueCommercetools", "")),
        "valueRoyalCyber": clean(d.get("valueRoyalCyber", "")),
    }

def _acc_wrap(items, meta):
    return {
        "org": meta.get("org") or ORG_DEFAULT,
        "title": meta.get("title") or ACC_TITLE_DEFAULT,
        "heading": meta.get("heading") or ACC_HEADING_DEFAULT,
        "intro": meta.get("intro") or ACC_INTRO_DEFAULT,
        "generated": meta.get("generated") or today(),
        "items": items,
    }

def accelerators_from_docx(path):
    from docx import Document
    doc = Document(path)
    items, meta, idx = [], {}, 1
    for t in doc.tables:
        d = {}
        for r in t.rows:
            cells = r.cells
            if len(cells) < 2:
                continue
            key = norm_key(cells[0].text)
            val = cells[1].text
            if key in META_KEYS:
                meta[META_KEYS[key]] = clean(val)
            elif key in ACC_FIELD_MAP:
                tgt = ACC_FIELD_MAP[key]
                d[tgt] = lines(val) if tgt in ("benefits", "technical", "resolves", "signals") else val
        if clean(d.get("name", "")):
            items.append(_build_item(d, idx)); idx += 1
    return _acc_wrap(items, meta)

def accelerators_from_xlsx(path):
    import openpyxl
    wb = openpyxl.load_workbook(path, data_only=True)
    meta = {}
    if "Meta" in wb.sheetnames:
        for r in wb["Meta"].iter_rows(values_only=True):
            if r and len(r) >= 2 and norm_key(r[0]) in META_KEYS:
                meta[META_KEYS[norm_key(r[0])]] = clean(r[1])
    ws = wb["Accelerators"] if "Accelerators" in wb.sheetnames else wb.active
    rows = [list(r) for r in ws.iter_rows(values_only=True)]
    rows = [r for r in rows if any(clean(c) for c in r)]
    if not rows:
        return _acc_wrap([], meta)
    header = [norm_key(c) for c in rows[0]]
    items, idx = [], 1
    for r in rows[1:]:
        d = {}
        for i, h in enumerate(header):
            if h in ACC_FIELD_MAP and i < len(r):
                tgt = ACC_FIELD_MAP[h]
                d[tgt] = lines(r[i]) if tgt in ("benefits", "technical", "resolves", "signals") else r[i]
        if clean(d.get("name", "")):
            items.append(_build_item(d, idx)); idx += 1
    return _acc_wrap(items, meta)

# ------------------------------------------------------------------------------- DRIVER
def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--kind", required=True, choices=["projects", "accelerators"])
    ap.add_argument("input")
    a = ap.parse_args()
    ext = os.path.splitext(a.input)[1].lower()
    if ext == ".json":
        with open(a.input, encoding="utf-8") as f:
            data = json.load(f)
    elif ext == ".docx":
        data = projects_from_docx(a.input) if a.kind == "projects" else accelerators_from_docx(a.input)
    elif ext in (".xlsx", ".xlsm"):
        data = projects_from_xlsx(a.input) if a.kind == "projects" else accelerators_from_xlsx(a.input)
    else:
        sys.exit(f"Unsupported extension: {ext}")
    json.dump(data, sys.stdout, ensure_ascii=False, indent=2)
    sys.stdout.write("\n")

if __name__ == "__main__":
    main()
